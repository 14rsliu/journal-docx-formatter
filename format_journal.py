#!/usr/bin/env python3
"""
Journal DOCX Formatter — Water Research / Elsevier style

Reformats a .docx file to match Water Research (Elsevier) submission format:
- Font: 12pt Times New Roman throughout (10pt for author info)
- Line spacing: single (1.0), 0.5 line space after
- Margins: 1 inch (2.54 cm) all sides
- Headings: bold, numbered (1. / 1.1 / 1.1.1), left-aligned
- Title: not bold
- Captions: only label bold (Fig. 1. / Table 1), no italic, centered
- Figures/tables: centered
- Body text: justified, no bold, no dashes
- Author email → footnote
- Equations: preserved with numbering
- References: page break before

Usage:
    python format_journal.py input.docx [--output output.docx]
"""

import argparse
import copy
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "Times New Roman"
FONT_COLOR = RGBColor(0, 0, 0)  # All text black

# Patterns that identify caption paragraphs
CAPTION_PATTERNS = [
    re.compile(r'^(Figure|Fig\.?|Table|Tab\.?)\s*\d', re.IGNORECASE),
    re.compile(r'^(图|表)\s*\d'),
]

# Regex to extract the label portion of a caption (e.g., "Fig. 1." or "Table 2")
CAPTION_LABEL_RE = re.compile(
    r'^((?:Figure|Fig\.?|Table|Tab\.?|图|表)\s*\d+[a-z]?\.?)\s*', re.IGNORECASE
)

# Patterns that identify section headings by text content (numbered sections)
HEADING_TEXT_PATTERNS = [
    re.compile(r'^\d+\.\s+\S'),          # Level 1: "1. Introduction"
    re.compile(r'^\d+\.\d+\s+\S'),       # Level 2: "2.1 Methods"
    re.compile(r'^\d+\.\d+\.\d+\s+\S'),  # Level 3: "2.1.1 Sub"
]

# Known section names (case-insensitive) for detecting headings without numbering
KNOWN_SECTIONS = [
    "abstract", "introduction", "methods", "methodology", "materials and methods",
    "results", "results and discussion", "discussion", "conclusion", "conclusions",
    "references", "acknowledgements", "acknowledgments", "appendix",
    "supplementary material", "supporting information",
    "credit authorship contribution statement",
    "declaration of competing interests",
    "declaration of generative ai and ai-assisted technologies in the writing process",
    "data availability",
]

# Section that requires a page break before it
PAGE_BREAK_BEFORE_SECTIONS = ["references"]

# Styles that indicate author/affiliation info (between Title and Abstract)
AUTHOR_STYLES = {"author", "subtitle"}

# Patterns for detecting author/affiliation text
AUTHOR_TEXT_PATTERNS = [
    re.compile(r'\*\s*corresponding\s+author', re.IGNORECASE),
    re.compile(r'@[\w.-]+\.\w{2,}'),  # email
    re.compile(r'^(department|school|college|faculty|institute|university|center|centre)\s+of\b', re.IGNORECASE),
]

# Email extraction pattern
EMAIL_RE = re.compile(r'[\w.-]+@[\w.-]+\.\w{2,}')

# Keywords pattern: "Keywords: word1 ,word2 ,word3"
KEYWORDS_RE = re.compile(r'^(Keywords?\s*:\s*)', re.IGNORECASE)


# ---------------------------------------------------------------------------
# Detection helpers
# ---------------------------------------------------------------------------

def is_author_info(para, before_abstract):
    """Check if paragraph is author/affiliation info (between Title and Abstract)."""
    if not before_abstract:
        return False
    style = (para.style.name if para.style else "").lower()
    if style in AUTHOR_STYLES:
        return True
    text = para.text.strip()
    if not text:
        return False
    for pat in AUTHOR_TEXT_PATTERNS:
        if pat.search(text):
            return True
    return False


def is_email_paragraph(para):
    """Check if paragraph contains corresponding author email info."""
    text = para.text.strip().lower()
    return bool(re.search(r'\*\s*corresponding\s+author', text, re.IGNORECASE)) or \
           (bool(EMAIL_RE.search(para.text)) and 'correspond' in text.lower())


def is_author_name_paragraph(para):
    """Check if paragraph is the author names line (comma-separated names with CAPS surnames)."""
    text = para.text.strip()
    if not text or len(text) > 300:
        return False
    # Pattern: "Firstname SURNAME, Firstname SURNAME, ..."
    # At least 2 comma-separated name segments with uppercase words
    parts = [p.strip() for p in text.split(',') if p.strip()]
    if len(parts) < 2:
        return False
    caps_count = sum(1 for p in parts if re.search(r'\b[A-Z]{2,}\b', p))
    return caps_count >= 2


def is_affiliation_paragraph(para):
    """Check if paragraph is an affiliation/institution line."""
    text = para.text.strip()
    if not text or len(text) > 300:
        return False
    keywords = ['university', 'department', 'institute', 'college', 'school of',
                'faculty', 'laboratory', 'center', 'centre']
    lower = text.lower()
    return any(k in lower for k in keywords)


def add_author_superscripts(doc):
    """Add superscript affiliation markers to author names and institution lines.

    Scans paragraphs between title and abstract. Collects affiliation paragraphs,
    assigns each a letter label (a, b, c...), then adds superscript markers after
    each author name and before each affiliation.

    If all authors share one affiliation, uses a single marker.
    """
    # Find title and abstract boundaries
    title_idx = abstract_idx = None
    for i, p in enumerate(doc.paragraphs):
        if is_title_paragraph(p) and title_idx is None:
            title_idx = i
        text = p.text.strip()
        if text == 'Abstract' or (p.style and p.style.name.lower() == 'abstract title'):
            abstract_idx = i
            break

    if title_idx is None or abstract_idx is None:
        return

    # Collect author and affiliation paragraphs between title and abstract
    author_paras = []
    affil_paras = []
    for i in range(title_idx + 1, abstract_idx):
        p = doc.paragraphs[i]
        if is_email_paragraph(p):
            continue
        if is_author_name_paragraph(p):
            author_paras.append(p)
        elif is_affiliation_paragraph(p):
            affil_paras.append(p)

    if not author_paras or not affil_paras:
        return

    # Check if superscripts already exist
    for p in author_paras:
        for run in p.runs:
            if run.font.superscript:
                return  # Already has superscripts

    # Assign labels: a, b, c, ...
    labels = [chr(ord('a') + j) for j in range(len(affil_paras))]

    # For single affiliation: all authors get same label
    if len(affil_paras) == 1:
        # Add superscript after each author name
        for p in author_paras:
            runs = list(p.runs)
            for run in runs:
                text = run.text
                if not text or not text.strip():
                    continue
                # Check if this run is a name (not just comma)
                if re.search(r'[A-Z]{2,}', text):
                    # Append superscript run after this one
                    sup_run = p.add_run(labels[0])
                    sup_run.font.superscript = True
                    sup_run.font.name = FONT_NAME
                    # Move the superscript run right after the current run in XML
                    run._element.addnext(sup_run._element)

        # Add superscript label before affiliation text
        for j, p in enumerate(affil_paras):
            if p.runs:
                first_run = p.runs[0]
                # Insert label at the beginning
                sup_run = p.add_run(labels[j])
                sup_run.font.superscript = True
                sup_run.font.name = FONT_NAME
                # Move before the first run
                first_run._element.addprevious(sup_run._element)
                # Add a space after the superscript
                first_run.text = ' ' + first_run.text.lstrip()
    else:
        # Multiple affiliations: need to map authors to affiliations
        # Default: assign all labels to all authors (user can adjust manually)
        all_labels = ','.join(labels)
        for p in author_paras:
            runs = list(p.runs)
            for run in runs:
                text = run.text
                if not text or not text.strip():
                    continue
                if re.search(r'[A-Z]{2,}', text):
                    sup_run = p.add_run(all_labels)
                    sup_run.font.superscript = True
                    sup_run.font.name = FONT_NAME
                    run._element.addnext(sup_run._element)

        for j, p in enumerate(affil_paras):
            if p.runs:
                first_run = p.runs[0]
                sup_run = p.add_run(labels[j])
                sup_run.font.superscript = True
                sup_run.font.name = FONT_NAME
                first_run._element.addprevious(sup_run._element)
                first_run.text = ' ' + first_run.text.lstrip()


def add_corresponding_author_star(doc):
    """Add * superscript after the corresponding author's name.

    Scans paragraphs between title and abstract. If a corresponding-author
    email line is found, extracts the author name from it (if present) or
    marks the last author in the author-names paragraph. The * is placed
    as a superscript run right after the matched name.
    """
    from docx.oxml.ns import qn

    title_idx = abstract_idx = None
    for i, p in enumerate(doc.paragraphs):
        if is_title_paragraph(p) and title_idx is None:
            title_idx = i
        text = p.text.strip()
        if text == 'Abstract' or (p.style and p.style.name.lower() == 'abstract title'):
            abstract_idx = i
            break

    if title_idx is None or abstract_idx is None:
        return

    # Find author names paragraph and email paragraph
    author_para = None
    email_para = None
    for i in range(title_idx + 1, abstract_idx):
        p = doc.paragraphs[i]
        if is_author_name_paragraph(p):
            author_para = p
        if is_email_paragraph(p):
            email_para = p

    if not author_para:
        return

    # Check if * superscript already exists
    for run in author_para.runs:
        if run.font.superscript and '*' in (run.text or ''):
            return  # Already has * superscript

    # Try to extract corresponding author name from email line
    # Patterns: "* Corresponding author: John SMITH (email@...)"
    #           "* Corresponding author. Email: email@..."
    target_name = None
    if email_para:
        email_text = email_para.text.strip()
        # Try to find a name with CAPS surname in the email line
        name_match = re.search(r':\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+[A-Z]{2,})', email_text)
        if name_match:
            target_name = name_match.group(1).strip()

    # Find the last name-like run in the author paragraph to add * after
    # If we have a target_name, try to match it; otherwise mark the last author
    last_name_run = None
    if target_name:
        # Search for the target name across runs
        full_text = author_para.text
        pos = full_text.find(target_name)
        if pos >= 0:
            # Find the run that contains the end of this name
            char_count = 0
            for run in author_para.runs:
                char_count += len(run.text or '')
                if char_count >= pos + len(target_name):
                    last_name_run = run
                    break

    if not last_name_run:
        # Default: add * after the last name-containing run
        for run in reversed(list(author_para.runs)):
            if run.text and re.search(r'[A-Z]{2,}', run.text):
                last_name_run = run
                break

    if last_name_run:
        sup_run = author_para.add_run('*')
        sup_run.font.superscript = True
        sup_run.font.name = FONT_NAME
        sup_run.font.size = Pt(11)
        last_name_run._element.addnext(sup_run._element)


def is_caption(text):
    """Check if paragraph text looks like a figure/table caption."""
    text = text.strip()
    for pat in CAPTION_PATTERNS:
        if pat.match(text):
            return True
    return False


def is_heading(para):
    """Check if paragraph is a heading (H1-H3) by style or text pattern."""
    style = para.style.name if para.style else ""
    if style.startswith("Heading"):
        try:
            level = int(style.replace("Heading", "").strip())
            return level if level <= 3 else 0
        except ValueError:
            pass

    text = para.text.strip()
    if not text:
        return 0
    for level, pat in enumerate(HEADING_TEXT_PATTERNS, start=1):
        if pat.match(text):
            return level

    text_lower = text.lower().rstrip('.')
    if len(text) < 100 and (text_lower in KNOWN_SECTIONS or
                            any(text_lower.startswith(s) for s in KNOWN_SECTIONS)):
        return 1

    return 0


def needs_page_break_before(text):
    """Check if this section should start on a new page."""
    text = text.strip().lower()
    for section in PAGE_BREAK_BEFORE_SECTIONS:
        if text == section or re.match(r'^\d+\.?\s*' + section + r'$', text, re.IGNORECASE):
            return True
    return False


def is_image_paragraph(para):
    """Check if paragraph contains an image."""
    for run in para.runs:
        if run._element.findall(qn("w:drawing")) or run._element.findall(qn("w:pict")):
            return True
    if para._element.findall(".//" + qn("w:drawing")):
        return True
    if para._element.findall(".//" + qn("w:pict")):
        return True
    return False


def is_equation_paragraph(para):
    """Check if paragraph contains OMML equation elements."""
    return bool(para._element.findall(".//" + qn("m:oMath")) or
                para._element.findall(".//" + qn("m:oMathPara")))


def is_title_paragraph(para):
    """Check if paragraph is the document title."""
    style = (para.style.name if para.style else "").lower()
    return style == "title"


def has_word_numbering(para):
    """Check if paragraph has Word built-in numbering (w:numPr in XML)."""
    numPr = para._element.find('.//' + qn('w:numPr'))
    if numPr is not None:
        numId = numPr.find(qn('w:numId'))
        # numId=0 means no numbering
        if numId is not None and numId.get(qn('w:val'), '0') != '0':
            return True
    return False


def is_bullet_paragraph(para):
    """Check if paragraph is a bullet/list item (by style, text pattern, or Word numbering).

    Detects: bullet symbols, dashes, numbered lists, Word built-in numbering (w:numPr).
    Does NOT match numbered section headings like '1. Introduction'.
    """
    style = (para.style.name if para.style else "").lower()
    if any(kw in style for kw in ("list", "bullet", "compact")):
        return True
    # Word built-in numbering (numbers not in text, generated by XML)
    if has_word_numbering(para):
        return True
    text = para.text.strip()
    if not text:
        return False
    # Bullet symbols
    if text[0] in ('\u2022', '•', '-', '\u2013', '\u2014'):
        return True
    # Numbered list: "1. something" or "1) something" — but NOT section headings
    m = re.match(r'^(\d+)[.)]\s+(.+)', text)
    if m:
        rest = m.group(2)
        rest_lower = rest.lower().rstrip('.')
        if rest_lower in KNOWN_SECTIONS:
            return False
        if re.match(r'^\d+\.\d+', text):
            return False
        return True
    # Lettered list: "(a) something" or "a) something"
    if re.match(r'^[(\[]?[a-z][)\]]\s+', text):
        return True
    return False


def is_caption_style(para):
    """Check if paragraph has a caption style (Image Caption, Table Caption, etc.)."""
    style = (para.style.name if para.style else "").lower()
    return "caption" in style


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def set_run_font(run, font_name=FONT_NAME, size=None, bold=None, italic=None, color=FONT_COLOR):
    """Set font properties on a run. Color defaults to black."""
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)

    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_alignment(para, alignment):
    """Set paragraph alignment."""
    para.paragraph_format.alignment = alignment


def replace_dashes_in_text(text, skip_number_ranges=False):
    """Replace en-dash and em-dash with appropriate alternatives.

    - En-dash between numbers (e.g. 15–30) → "15 to 30"
    - Hyphen between numbers (e.g. 15-30) → "15 to 30"
    - Em-dash / en-dash as punctuation → comma or space

    Args:
        skip_number_ranges: If True, do NOT convert number-dash-number to "to".
            Used for references, DOIs, page ranges, etc.
    """
    if not skip_number_ranges:
        # En-dash between numbers: 15–30 → 15 to 30
        text = re.sub(r'(\d+\.?\d*%?)\s*\u2013\s*(\d)', r'\1 to \2', text)
        # Hyphen between numbers: 15-30 → 15 to 30 (but NOT compound words like DP-SGD)
        text = re.sub(r'(\d+\.?\d*%?)\s*-\s*(\d)', r'\1 to \2', text)
    else:
        # Still convert en-dash between numbers to hyphen for consistency
        text = re.sub(r'(\d)\u2013(\d)', r'\1-\2', text)
    # Em-dash or en-dash as sentence punctuation → comma
    text = text.replace('\u2014', ', ')   # em-dash
    text = text.replace('\u2013', ', ')   # remaining en-dash (non-numeric)
    # Clean up double spaces
    text = re.sub(r'  +', ' ', text)
    # Clean up ", ," or ", ,"
    text = re.sub(r',\s*,', ',', text)
    return text


def replace_tildes_in_text(text):
    """Replace tilde (~) and tilde operator (∼ U+223C) with 'approximately' or remove.

    - ~/∼ + number → approximately number  (e.g. ~50 → approximately 50)
    - other ~/∼ → remove
    """
    # ~/∼ + number → approximately number
    text = re.sub(r'[~∼]\s*(\d)', r'approximately \1', text)
    # Remove remaining tildes and tilde operators
    text = text.replace('~', '')
    text = text.replace('∼', '')
    # Clean up double spaces
    text = re.sub(r'  +', ' ', text)
    return text


def replace_math_symbols_in_text(text):
    """Replace mathematical symbols with natural language in body text.

    Handles: ≈, ≥, ≤, <, > (before numbers), ± (in context).
    Skips table cells and equations (caller responsibility).
    """
    # ≈ number → approximately number (ensure space before when preceded by non-space char)
    text = re.sub(r'(?<=\S)≈\s*(\d)', r' approximately \1', text)
    text = re.sub(r'≈\s*(\d)', r'approximately \1', text)
    # remaining ≈ → approximately (with space guard)
    text = re.sub(r'(?<=\S)≈', ' approximately ', text)
    text = text.replace('≈', 'approximately ')

    # ± between numbers is standard academic notation (mean ± std) — preserve it
    # Only remove orphan ± not adjacent to numbers
    text = re.sub(r'(?<!\d[\s%])±(?!\s*\d)', '', text)

    # ≥ number → at least number (ensure space before when preceded by non-space char)
    text = re.sub(r'(?<=\S)≥\s*(\d)', r' at least \1', text)
    text = re.sub(r'≥\s*(\d)', r'at least \1', text)
    # ≤ number → no more than number (ensure space before when preceded by non-space char)
    text = re.sub(r'(?<=\S)≤\s*(\d)', r' no more than \1', text)
    text = re.sub(r'≤\s*(\d)', r'no more than \1', text)

    # > number (not inside HTML tags) → greater than number
    text = re.sub(r'(?<![<\w])>\s*(\d)', r'greater than \1', text)
    # < number → less than number
    text = re.sub(r'(?<!\w)<\s*(\d)', r'less than \1', text)

    # Clean up double spaces
    text = re.sub(r'  +', ' ', text)
    return text


def reduce_relative_improvements(text):
    """Remove relative improvement phrases like 'by X percentage points' or 'within X pp of Y'.

    Keeps absolute values (e.g. '90.46% accuracy') but removes comparative deltas.
    Examples removed:
      - ', within 0.26 percentage points of centralized training'
      - ', outperforming local-only training by approximately 12 percentage points'
      - 'by 3.02 percentage points'
      - 'exceeding full participation by 3.02 percentage points'
    """
    # "within X (percentage points|pp) of ..." — remove the clause
    text = re.sub(
        r',?\s*within\s+[\d.]+\s+(?:percentage\s+points?|pp)\s+of\s+[^.,;]+',
        '', text, flags=re.IGNORECASE)
    # "within X pp" without "of" (e.g. "matches centralized accuracy within 0.26 pp")
    text = re.sub(
        r'\s+within\s+[\d.]+\s+(?:percentage\s+points?|pp)\b',
        '', text, flags=re.IGNORECASE)
    # "outperforming/exceeding/surpassing ... by (approximately) X percentage points"
    text = re.sub(
        r',?\s*(?:outperforming|exceeding|surpassing|improving\s+over)\s+[^.,;]*?by\s+(?:approximately\s+)?[\d.]+\s+(?:percentage\s+points?|pp)',
        '', text, flags=re.IGNORECASE)
    # "by (approximately) X percentage points" (standalone, at end of clause)
    text = re.sub(
        r',?\s*by\s+(?:approximately\s+)?[\d.]+\s+(?:percentage\s+points?|pp)\b',
        '', text, flags=re.IGNORECASE)
    # "X pp higher/lower/more/less than ..."
    text = re.sub(
        r'[\d.]+\s+(?:percentage\s+points?|pp)\s+(?:higher|lower|more|less|above|below)\s+(?:than\s+)?[^.,;]*',
        '', text, flags=re.IGNORECASE)
    # Clean up artifacts: double commas, leading commas, double spaces
    text = re.sub(r',\s*,', ',', text)
    text = re.sub(r'\.\s*,', '.', text)
    text = re.sub(r',\s*\.', '.', text)
    text = re.sub(r'  +', ' ', text)
    text = text.strip()
    return text


def reduce_parentheses(text):
    """Remove or unwrap parenthetical content to reduce bracket usage.

    Strategy:
    - Short parentheticals (<=30 chars) that are just abbreviations/refs: keep
    - Parentheticals with full clauses: unwrap (remove parens, integrate text)
    - Nested numeric ranges like (84.82%-89.12%): keep as-is
    """
    def _handle_paren(m):
        content = m.group(1).strip()
        # Keep citations like (Smith et al., 2020) or [1]
        if re.match(r'^[\w\s.,&]+\d{4}', content):
            return m.group(0)
        # Keep short abbreviations/acronyms like (NRW), (FL), (PBI)
        if len(content) <= 8 and re.match(r'^[A-Z0-9./\s]+$', content):
            return m.group(0)
        # Keep numeric ranges like (84.82%-89.12%)
        if re.match(r'^[\d.%\-\u2013 ]+$', content):
            return m.group(0)
        # Keep single values like (90.46%)
        if re.match(r'^[\d.]+%?$', content):
            return m.group(0)
        # Keep "e.g., ..." or "i.e., ..."
        if content.lower().startswith(('e.g.', 'i.e.')):
            return m.group(0)
        # For longer content: unwrap — replace parens with commas
        # Check if preceded by comma/period already
        return ', ' + content + ','

    text = re.sub(r'\(([^()]+)\)', _handle_paren, text)
    # Clean up: ", ," → ","
    text = re.sub(r',\s*,', ',', text)
    # Clean up leading comma after period: ". , " → ". "
    text = re.sub(r'\.\s*,\s*', '. ', text)
    text = re.sub(r'  +', ' ', text)
    return text


def format_highlight_bullet(para):
    """Add bullet '• ' prefix to a highlight paragraph if not already present."""
    text = para.text.strip()
    if not text or text.startswith('\u2022') or text.startswith('•'):
        return
    # Prepend bullet to the first run
    if para.runs:
        para.runs[0].text = '\u2022 ' + para.runs[0].text


def format_keywords_paragraph(para, body_size):
    """Format keywords: bold 'Keywords:' label, fix spacing (', keyword' not ' ,keyword')."""
    full_text = para.text
    m = KEYWORDS_RE.match(full_text)
    if not m:
        return

    label_end = m.end(1)

    # Fix keyword spacing in runs: " ,word" → ", word"
    for run in para.runs:
        if run.text:
            # Fix pattern: "space comma word" or "space comma space word"
            run.text = re.sub(r'\s+,\s*', ', ', run.text)
            # Also ensure no trailing space after last keyword
            # (will be handled by overall cleanup)

    # Now bold only the label portion
    pos = 0
    for run in para.runs:
        run_len = len(run.text)
        run_end = pos + run_len

        if run_end <= label_end:
            set_run_font(run, size=body_size, bold=True, italic=False, color=FONT_COLOR)
        elif pos >= label_end:
            set_run_font(run, size=body_size, bold=False, italic=False, color=FONT_COLOR)
        else:
            # Split at boundary
            label_part = run.text[:label_end - pos]
            rest_part = run.text[label_end - pos:]
            run.text = label_part
            set_run_font(run, size=body_size, bold=True, italic=False, color=FONT_COLOR)

            new_r = copy.deepcopy(run._element)
            run._element.addnext(new_r)
            from docx.text.run import Run
            new_run = Run(new_r, run._parent)
            new_run.text = rest_part
            set_run_font(new_run, size=body_size, bold=False, italic=False, color=FONT_COLOR)

        pos = run_end


def format_caption_runs(para, body_size):
    """Format caption: only label (Fig. 1. / Table 1) is bold, rest is normal. No italic. Centered."""
    full_text = para.text
    m = CAPTION_LABEL_RE.match(full_text)
    if not m:
        # Fallback: format all runs as normal
        for run in para.runs:
            set_run_font(run, size=body_size, bold=False, italic=False, color=FONT_COLOR)
        return

    label_end = m.end(1)  # end position of "Fig. 1" part

    # Track position across runs and set bold/not-bold accordingly
    pos = 0
    for run in para.runs:
        run_len = len(run.text)
        run_end = pos + run_len

        if run_end <= label_end:
            # Entire run is within label → bold
            set_run_font(run, size=body_size, bold=True, italic=False, color=FONT_COLOR)
        elif pos >= label_end:
            # Entire run is after label → not bold
            set_run_font(run, size=body_size, bold=False, italic=False, color=FONT_COLOR)
        else:
            # Run spans the boundary — need to split
            # Text before boundary
            label_part = run.text[:label_end - pos]
            desc_part = run.text[label_end - pos:]
            run.text = label_part
            set_run_font(run, size=body_size, bold=True, italic=False, color=FONT_COLOR)

            # Insert new run after current one for the description part
            new_r = copy.deepcopy(run._element)
            run._element.addnext(new_r)
            from docx.text.run import Run
            new_run = Run(new_r, run._parent)
            new_run.text = desc_part
            set_run_font(new_run, size=body_size, bold=False, italic=False, color=FONT_COLOR)

        pos = run_end


def _set_cell_border(cell, border_name, sz="4", val="single", color="000000"):
    """Set a specific border on a table cell (top, bottom, left, right, etc.)."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    border = tcBorders.find(qn(f"w:{border_name}"))
    if border is None:
        border = OxmlElement(f"w:{border_name}")
        tcBorders.append(border)
    border.set(qn("w:val"), val)
    border.set(qn("w:sz"), sz)
    border.set(qn("w:color"), color)
    border.set(qn("w:space"), "0")


def _clear_cell_borders(cell):
    """Remove all borders from a cell."""
    for side in ("top", "bottom", "left", "right", "insideH", "insideV"):
        _set_cell_border(cell, side, val="none", sz="0")


def format_table_fonts(table, font_name=FONT_NAME, size=None):
    """Format table: three-line style (三线表), full width, even columns, Times New Roman."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Center the table
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tblPr.append(jc)
    jc.set(qn("w:val"), "center")

    # Set table width to 100% (5000 fifths of a percent = 100%)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    # Remove table-level borders first (we'll set cell-level for three-line)
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Fixed layout so column widths are respected exactly
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

    rows = list(table.rows)
    num_rows = len(rows)
    num_cols = len(table.columns) if table.columns else 0

    # Content-aware column width distribution to minimize table height.
    # Measures actual content length (incl. oMath), skips merged cells,
    # computes ideal single-line widths, then scales to fill page width.
    # Page usable width: 9360 twips (8.5" - 2*1" margins = 6.5" * 1440 twips/in)
    PAGE_WIDTH_TWIPS = 9360
    # Times New Roman 11pt: ~120 twips per average char; cell padding ~180 twips
    CHAR_WIDTH = 120
    CELL_PADDING = 180
    MIN_COL_WIDTH = 400  # minimum ~0.28 inches

    if num_cols > 0:
        # Measure max content char count per column (including oMath text)
        col_max_chars = [0] * num_cols
        for row in rows:
            # Detect merged row: all cells share same _tc
            cell_ids = [id(c._tc) for c in row.cells]
            if len(set(cell_ids)) == 1 and num_cols > 1:
                continue  # skip fully-merged rows
            for idx, cell in enumerate(row.cells):
                if idx >= num_cols:
                    break
                tc = cell._tc
                # Gather all text including oMath
                all_text = ''
                for p_el in tc.findall('.//' + qn('w:p')):
                    for r_el in p_el.findall(qn('w:r')):
                        t_el = r_el.find(qn('w:t'))
                        if t_el is not None and t_el.text:
                            all_text += t_el.text
                    for mt in p_el.findall('.//' + qn('m:t')):
                        if mt.text:
                            all_text += mt.text
                char_count = len(all_text.strip())
                col_max_chars[idx] = max(col_max_chars[idx], char_count)

        # Compute ideal width for each column (single-line display)
        col_ideal = [max(c * CHAR_WIDTH + CELL_PADDING, MIN_COL_WIDTH)
                     for c in col_max_chars]
        total_ideal = sum(col_ideal)

        if total_ideal <= PAGE_WIDTH_TWIPS:
            # All columns fit; distribute remaining space proportionally
            remaining = PAGE_WIDTH_TWIPS - total_ideal
            col_twips = [int(w + remaining * w / total_ideal) for w in col_ideal]
        else:
            # Need to compress; scale down proportionally
            col_twips = [max(int(PAGE_WIDTH_TWIPS * w / total_ideal), MIN_COL_WIDTH)
                         for w in col_ideal]

        # Fix rounding: last col gets remainder to sum exactly to PAGE_WIDTH_TWIPS
        col_twips[-1] = PAGE_WIDTH_TWIPS - sum(col_twips[:-1])

        # Update tblGrid/gridCol to match
        tblGrid = tbl.find(qn("w:tblGrid"))
        if tblGrid is None:
            tblGrid = OxmlElement("w:tblGrid")
            tbl.insert(1, tblGrid)  # after tblPr
        else:
            for old_gc in tblGrid.findall(qn("w:gridCol")):
                tblGrid.remove(old_gc)
        for tw in col_twips:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(tw))
            tblGrid.append(gc)

        # Update each cell width (dxa type = twips)
        for row in rows:
            for idx, cell in enumerate(row.cells):
                if idx < num_cols:
                    tc = cell._tc
                    tcPr = tc.find(qn("w:tcPr"))
                    if tcPr is None:
                        tcPr = OxmlElement("w:tcPr")
                        tc.insert(0, tcPr)
                    tcW = tcPr.find(qn("w:tcW"))
                    if tcW is None:
                        tcW = OxmlElement("w:tcW")
                        tcPr.append(tcW)
                    tcW.set(qn("w:w"), str(col_twips[idx]))
                    tcW.set(qn("w:type"), "dxa")

    # Three-line table (三线表): top of header, bottom of header, bottom of table
    for row_idx, row in enumerate(rows):
        for cell in row.cells:
            _clear_cell_borders(cell)
            if row_idx == 0:
                # Header row: top border (thick) + bottom border (thin)
                _set_cell_border(cell, "top", sz="12", val="single", color="000000")
                _set_cell_border(cell, "bottom", sz="6", val="single", color="000000")
            elif row_idx == num_rows - 1:
                # Last row: bottom border (thick)
                _set_cell_border(cell, "bottom", sz="12", val="single", color="000000")
            # Middle rows: no borders (already cleared)

    # Compact table: minimize cell margins (top/bottom=0, left/right=small)
    tblCellMar = tblPr.find(qn("w:tblCellMar"))
    if tblCellMar is not None:
        tblPr.remove(tblCellMar)
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, val in [("top", "0"), ("bottom", "0"), ("left", "57"), ("right", "57")]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")
        tblCellMar.append(el)
    tblPr.append(tblCellMar)

    # Format cell text: Times New Roman, centered, no bold/italic, minimize wrapping
    for row_idx, row in enumerate(rows):
        # Set exact row height for compact layout
        tr = row._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            tr.insert(0, trPr)
        trHeight = trPr.find(qn("w:trHeight"))
        if trHeight is None:
            trHeight = OxmlElement("w:trHeight")
            trPr.append(trHeight)
        trHeight.set(qn("w:val"), "288")   # ~0.2in minimum row height
        trHeight.set(qn("w:hRule"), "atLeast")

        for cell in row.cells:
            # Set no-wrap on cell to minimize line breaks
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            noWrap = tcPr.find(qn("w:noWrap"))
            if noWrap is None:
                noWrap = OxmlElement("w:noWrap")
                tcPr.append(noWrap)
            # Vertical alignment: center
            vAlign = tcPr.find(qn("w:vAlign"))
            if vAlign is None:
                vAlign = OxmlElement("w:vAlign")
                tcPr.append(vAlign)
            vAlign.set(qn("w:val"), "center")

            for para in cell.paragraphs:
                set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.CENTER)
                # Zero paragraph spacing for compact layout
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                pPr = para._element.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    para._element.insert(0, pPr)
                # Single line spacing (240 twips = single)
                spacing = pPr.find(qn("w:spacing"))
                if spacing is None:
                    spacing = OxmlElement("w:spacing")
                    pPr.append(spacing)
                spacing.set(qn("w:line"), "240")
                spacing.set(qn("w:lineRule"), "auto")
                for run in para.runs:
                    is_header = (row_idx == 0)
                    set_run_font(run, font_name, size=size, bold=is_header, italic=False, color=FONT_COLOR)


def set_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set page margins in inches for all sections."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


# ---------------------------------------------------------------------------
# Footnote support (low-level XML)
# ---------------------------------------------------------------------------

def get_or_create_footnotes_part(doc):
    """Get or create the footnotes part in the document."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    # Check if footnotes part exists
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            return rel.target_part
    # Need to create — return None to signal caller
    return None


def add_footnote_via_xml(doc, para, footnote_text, font_size=10, custom_mark=None):
    """Add a footnote to a paragraph using low-level XML.

    Since python-docx has limited footnote support, we manipulate XML directly.
    This creates a footnote reference in the paragraph and the footnote content.

    Args:
        custom_mark: If set (e.g. "*"), use this symbol instead of auto-numbering.
            The symbol appears both in the body (as superscript) and in the
            footnote area (instead of the auto-number).
    """
    from lxml import etree

    # Access the footnotes part via blob
    footnotes_part = None
    footnotes_tree = None

    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            footnotes_part = rel.target_part
            footnotes_tree = etree.fromstring(footnotes_part.blob)
            break

    if footnotes_tree is None:
        return _add_footnote_as_paragraph(doc, para, footnote_text, font_size)

    # Namespace
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Find the next footnote ID
    existing = footnotes_tree.findall(f"{{{w}}}footnote")
    max_id = 0
    for fn in existing:
        fid = int(fn.get(f"{{{w}}}id", "0"))
        if fid > max_id:
            max_id = fid
    new_id = max_id + 1

    # Build footnote XML
    footnote = etree.SubElement(footnotes_tree, f"{{{w}}}footnote")
    footnote.set(f"{{{w}}}id", str(new_id))

    fn_para = etree.SubElement(footnote, f"{{{w}}}p")

    if custom_mark:
        # Custom mark: show the symbol instead of auto-number in footnote area
        fn_mark_run = etree.SubElement(fn_para, f"{{{w}}}r")
        fn_mark_rPr = etree.SubElement(fn_mark_run, f"{{{w}}}rPr")
        fn_mark_style = etree.SubElement(fn_mark_rPr, f"{{{w}}}rStyle")
        fn_mark_style.set(f"{{{w}}}val", "FootnoteReference")
        fn_mark_t = etree.SubElement(fn_mark_run, f"{{{w}}}t")
        fn_mark_t.text = custom_mark
    else:
        # Auto-numbered footnote ref inside footnote
        fn_ref_run = etree.SubElement(fn_para, f"{{{w}}}r")
        fn_ref_rPr = etree.SubElement(fn_ref_run, f"{{{w}}}rPr")
        fn_ref_style = etree.SubElement(fn_ref_rPr, f"{{{w}}}rStyle")
        fn_ref_style.set(f"{{{w}}}val", "FootnoteReference")
        etree.SubElement(fn_ref_run, f"{{{w}}}footnoteRef")

    # Space
    space_run = etree.SubElement(fn_para, f"{{{w}}}r")
    space_t = etree.SubElement(space_run, f"{{{w}}}t")
    space_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    space_t.text = " "

    # Footnote text run
    text_run = etree.SubElement(fn_para, f"{{{w}}}r")
    text_rPr = etree.SubElement(text_run, f"{{{w}}}rPr")
    text_rFonts = etree.SubElement(text_rPr, f"{{{w}}}rFonts")
    text_rFonts.set(f"{{{w}}}ascii", FONT_NAME)
    text_rFonts.set(f"{{{w}}}hAnsi", FONT_NAME)
    text_rFonts.set(f"{{{w}}}eastAsia", FONT_NAME)
    text_sz = etree.SubElement(text_rPr, f"{{{w}}}sz")
    text_sz.set(f"{{{w}}}val", str(font_size * 2))
    text_szCs = etree.SubElement(text_rPr, f"{{{w}}}szCs")
    text_szCs.set(f"{{{w}}}val", str(font_size * 2))
    text_t = etree.SubElement(text_run, f"{{{w}}}t")
    text_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_t.text = footnote_text

    # Write back the modified XML to the part blob
    footnotes_part._blob = etree.tostring(footnotes_tree, xml_declaration=True,
                                           encoding="UTF-8", standalone=True)

    if custom_mark:
        # Custom mark in body: use w:customMarkFollows + plain superscript text
        # instead of w:footnoteReference (which shows auto-number)
        ref_run = OxmlElement("w:r")
        ref_rPr = OxmlElement("w:rPr")
        ref_vertAlign = OxmlElement("w:vertAlign")
        ref_vertAlign.set(qn("w:val"), "superscript")
        ref_rPr.append(ref_vertAlign)
        # Set font
        ref_rFonts = OxmlElement("w:rFonts")
        ref_rFonts.set(qn("w:ascii"), FONT_NAME)
        ref_rFonts.set(qn("w:hAnsi"), FONT_NAME)
        ref_rPr.append(ref_rFonts)
        ref_run.append(ref_rPr)
        # The footnoteReference with customMarkFollows suppresses the auto-number
        ref_mark = OxmlElement("w:footnoteReference")
        ref_mark.set(qn("w:id"), str(new_id))
        ref_mark.set(qn("w:customMarkFollows"), "1")
        ref_run.append(ref_mark)
        # Add the custom symbol as text right after
        ref_t = OxmlElement("w:t")
        ref_t.text = custom_mark
        ref_run.append(ref_t)
        para._element.append(ref_run)
    else:
        # Standard auto-numbered footnote reference in the main paragraph
        ref_run = OxmlElement("w:r")
        ref_rPr = OxmlElement("w:rPr")
        ref_style = OxmlElement("w:rStyle")
        ref_style.set(qn("w:val"), "FootnoteReference")
        ref_rPr.append(ref_style)
        ref_run.append(ref_rPr)
        ref_mark = OxmlElement("w:footnoteReference")
        ref_mark.set(qn("w:id"), str(new_id))
        ref_run.append(ref_mark)
        para._element.append(ref_run)

    return True


def _add_footnote_as_paragraph(doc, para, footnote_text, font_size):
    """Fallback: add footnote text as a small paragraph after the given paragraph."""
    # Create a new paragraph element
    new_p = OxmlElement("w:p")

    # Paragraph properties - centered
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    new_p.append(pPr)

    # Text run
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_size * 2))
    rPr.append(szCs)
    # Color black
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), "000000")
    rPr.append(color_elem)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = footnote_text
    r.append(t)
    new_p.append(r)

    # Insert after the given paragraph
    para._element.addnext(new_p)
    return True


# ---------------------------------------------------------------------------
# Main formatting
# ---------------------------------------------------------------------------

def merge_bullets_to_prose(doc, section_name="introduction"):
    """Merge bullet/list paragraphs within a section into flowing prose.

    Bullets are joined with '; ' and appended to the preceding normal paragraph,
    or combined into a single new paragraph if there is no preceding text.
    """
    paras = list(doc.paragraphs)
    current_section = None
    i = 0
    while i < len(paras):
        p = paras[i]
        heading_level = is_heading(p)
        if heading_level == 1:
            sec_text = p.text.strip().lower()
            # Strip leading number: "1. Introduction" -> "introduction"
            sec_text = re.sub(r'^\d+\.?\s*', '', sec_text).rstrip('.')
            current_section = sec_text

        if current_section and section_name in current_section and not heading_level and is_bullet_paragraph(p):
            # Collect consecutive bullets
            bullet_texts = []
            start_idx = i
            while i < len(paras) and is_bullet_paragraph(paras[i]):
                # Remove Word numbering XML first
                if has_word_numbering(paras[i]):
                    remove_word_numbering(paras[i])
                bt = paras[i].text.strip()
                # Strip leading bullet symbols, numbers, letters
                bt = re.sub(r'^[\u2022\u2023\u25E6\u2043\u2022\-\u2013\u2014]\s*', '', bt)
                bt = re.sub(r'^\d+[.)]\s+', '', bt)
                bt = re.sub(r'^[(\[]?[a-z][)\]]\s+', '', bt)
                if bt:
                    # Lowercase first char for flow (unless acronym / proper noun)
                    if bt[0].isupper() and (len(bt) < 2 or not bt[1].isupper()):
                        bt = bt[0].lower() + bt[1:]
                    # Strip trailing period for joining
                    bt = bt.rstrip('.')
                    bullet_texts.append(bt)
                i += 1

            if bullet_texts:
                merged = '; '.join(bullet_texts) + '.'
                # Remove original bullet paragraphs (keep first, clear rest)
                first_bullet = paras[start_idx]
                # Clear all runs and set merged text
                for run in first_bullet.runs:
                    run.text = ""
                if first_bullet.runs:
                    first_bullet.runs[0].text = merged
                else:
                    first_bullet.add_run(merged)
                # Delete subsequent bullet paragraphs
                for j in range(start_idx + 1, start_idx + len(bullet_texts)):
                    if j < len(paras):
                        p_del = paras[j]._element
                        p_del.getparent().remove(p_del)
                # Re-read paras after deletion
                paras = list(doc.paragraphs)
                i = start_idx + 1
        else:
            i += 1


def remove_word_numbering(para):
    """Remove Word built-in numbering (w:numPr) from a paragraph's XML."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)


def normalize_bullet_symbol(para):
    """Replace dash/number/letter bullet prefixes with circular bullet '\\u2022'.
    Also removes Word built-in numbering and adds bullet prefix."""
    # Handle Word XML numbering (numbers not in text)
    if has_word_numbering(para):
        remove_word_numbering(para)
        # Add bullet prefix to text since the number was auto-generated
        if para.runs:
            text = para.runs[0].text or ""
            if not text.startswith('\u2022'):
                para.runs[0].text = '\u2022 ' + text
        return

    first_run = True
    for run in para.runs:
        if run.text and first_run:
            # Replace leading dash/en-dash/em-dash
            run.text = re.sub(r'^[\-\u2013\u2014]\s*', '\u2022 ', run.text, count=1)
            # Replace leading numbered list: "1. " "2) " etc.
            run.text = re.sub(r'^\d+[.)]\s+', '\u2022 ', run.text, count=1)
            # Replace leading lettered list: "(a) " "a) " "[a] " etc.
            run.text = re.sub(r'^[(\[]?[a-z][)\]]\s+', '\u2022 ', run.text, count=1)
            first_run = False


def format_document(docx_path, output_path=None,
                    body_size=11, author_size=11, table_size=11,
                    line_spacing=1.0, space_after=0.5,
                    margin=1.0):
    """
    Apply Water Research / Elsevier journal formatting to a .docx file.
    """
    doc = Document(docx_path)

    # Set page margins
    set_margins(doc, top=margin, bottom=margin, left=margin, right=margin)

    # --- Pass 1: Extract email from corresponding author and convert to footnote ---
    email_para = None
    author_name_para = None
    before_abstract = False
    title_seen = False

    for para in doc.paragraphs:
        style_name = (para.style.name if para.style else "").lower()
        text = para.text.strip()

        if style_name == "title":
            title_seen = True
            before_abstract = True
            continue
        if style_name in ("abstract title", "abstract") or text.lower() == "abstract":
            before_abstract = False
            break

        if before_abstract:
            if style_name in AUTHOR_STYLES:
                author_name_para = para
            if text and re.search(r'\*\s*corresponding\s+author', text, re.IGNORECASE):
                email_para = para

    # Move email to footnote
    if email_para and author_name_para:
        email_text = email_para.text.strip()
        # Strip leading * from email text since it will be the footnote custom mark
        email_text = re.sub(r'^\*\s*', '', email_text)
        add_footnote_via_xml(doc, author_name_para, email_text, font_size=9,
                            custom_mark="*")
        for run in email_para.runs:
            run.text = ""
        email_para.text = ""

    # --- Pass 1.4a: Add * superscript for corresponding author ---
    # Note: if email was moved to footnote with custom_mark="*", the * is already
    # added by the footnote reference. Only call this if no footnote was created.
    if not (email_para and author_name_para):
        add_corresponding_author_star(doc)

    # --- Pass 1.4b: Add superscript affiliation markers to author names ---
    add_author_superscripts(doc)

    # --- Pass 1.5: Merge bullet lists in Introduction into prose ---
    merge_bullets_to_prose(doc, section_name="introduction")

    # --- Pass 1.6: Process all oMath elements (paragraphs + tables) ---
    # Must run BEFORE Pass 2 so flattened text gets cleaned by subsequent passes.
    body_para_elements = set(id(p._element) for p in doc.paragraphs)
    all_para_elements = [p._element for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_para_elements.append(p._element)

    eq_counter = 0  # Counter for display equation numbering

    for p_elem in all_para_elements:
        is_in_table = id(p_elem) not in body_para_elements

        for omath in list(p_elem.findall('.//' + qn('m:oMath'))):
            mt_elements = omath.findall('.//' + qn('m:t'))
            mt_texts = [t.text for t in mt_elements if t.text]
            combined = ''.join(mt_texts).strip()

            parent = omath.getparent()
            parent_tag = parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag
            # Only treat as display equation if in body (not table) and inside oMathPara
            is_display = (parent_tag == 'oMathPara') and not is_in_table

            has_complex = (omath.findall('.//' + qn('m:f')) or
                           omath.findall('.//' + qn('m:rad')) or
                           omath.findall('.//' + qn('m:sSup')) or
                           omath.findall('.//' + qn('m:sSub')) or
                           omath.findall('.//' + qn('m:nary')))
            is_simple = not has_complex and len(combined) <= 30

            # For table oMathPara: treat as simple regardless (flatten to plain text)
            is_table_display = (parent_tag == 'oMathPara') and is_in_table
            should_flatten = (is_simple and combined and not is_display) or is_table_display

            if should_flatten and combined:
                # Flatten simple inline/table oMath to plain text run
                text = combined
                if '~' in text:
                    text = replace_tildes_in_text(text)
                if '\u2013' in text or '\u2014' in text:
                    text = replace_dashes_in_text(text)
                new_r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), FONT_NAME)
                rFonts.set(qn('w:hAnsi'), FONT_NAME)
                rFonts.set(qn('w:eastAsia'), FONT_NAME)
                rFonts.set(qn('w:cs'), FONT_NAME)
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(body_size * 2)))
                rPr.append(sz)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), str(int(body_size * 2)))
                rPr.append(szCs)
                color_el = OxmlElement('w:color')
                color_el.set(qn('w:val'), '000000')
                rPr.append(color_el)
                i_el = OxmlElement('w:i')
                i_el.set(qn('w:val'), '0')
                rPr.append(i_el)
                new_r.append(rPr)
                t_el = OxmlElement('w:t')
                t_el.set(qn('xml:space'), 'preserve')
                t_el.text = text
                new_r.append(t_el)
                if is_table_display:
                    # Table oMathPara: insert new_r before oMathPara at paragraph level, remove oMathPara
                    omath_para = parent  # parent is oMathPara
                    omath_para.addprevious(new_r)
                    omath_para.getparent().remove(omath_para)
                else:
                    omath.addprevious(new_r)
                    omath.getparent().remove(omath)
            else:
                # Complex or display oMath: set font, size, italic style
                for mr in omath.findall('.//' + qn('m:r')):
                    wrPr = mr.find(qn('w:rPr'))
                    if wrPr is None:
                        wrPr = OxmlElement('w:rPr')
                        mr.insert(0, wrPr)
                    rFonts = wrPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        wrPr.insert(0, rFonts)
                    rFonts.set(qn('w:ascii'), FONT_NAME)
                    rFonts.set(qn('w:hAnsi'), FONT_NAME)
                    rFonts.set(qn('w:eastAsia'), FONT_NAME)
                    rFonts.set(qn('w:cs'), FONT_NAME)
                    sz = wrPr.find(qn('w:sz'))
                    if sz is None:
                        sz = OxmlElement('w:sz')
                        wrPr.append(sz)
                    sz.set(qn('w:val'), str(int(body_size * 2)))
                    szCs = wrPr.find(qn('w:szCs'))
                    if szCs is None:
                        szCs = OxmlElement('w:szCs')
                        wrPr.append(szCs)
                    szCs.set(qn('w:val'), str(int(body_size * 2)))
                    color_el = wrPr.find(qn('w:color'))
                    if color_el is None:
                        color_el = OxmlElement('w:color')
                        wrPr.append(color_el)
                    color_el.set(qn('w:val'), '000000')
                    mrPr = mr.find(qn('m:rPr'))
                    if mrPr is None:
                        mrPr = OxmlElement('m:rPr')
                        mr.insert(0, mrPr)
                    sty = mrPr.find(qn('m:sty'))
                    if sty is None:
                        sty = OxmlElement('m:sty')
                        mrPr.append(sty)
                    if is_display:
                        sty.set(qn('m:val'), 'i')
                    else:
                        sty.set(qn('m:val'), 'p')

                # Clean text inside complex oMath
                for mt in omath.findall('.//' + qn('m:t')):
                    if mt.text:
                        if '~' in mt.text:
                            mt.text = replace_tildes_in_text(mt.text)
                        if '\u2013' in mt.text or '\u2014' in mt.text:
                            mt.text = replace_dashes_in_text(mt.text)

                # Add equation number to display equations: centered eq + right-aligned "(N)"
                # Strategy: unwrap oMath from oMathPara so we can use tab stops for layout
                # Result: [tab-center] [oMath] [tab-right] [(N)]
                if is_display:
                    eq_counter += 1

                    # Set paragraph alignment to left so tab stops control centering
                    pPr = p_elem.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        p_elem.insert(0, pPr)
                    jc = pPr.find(qn('w:jc'))
                    if jc is None:
                        jc = OxmlElement('w:jc')
                        pPr.append(jc)
                    jc.set(qn('w:val'), 'left')

                    # Set up tab stops: center at page center, right at right margin
                    # Page width 12240 twips (8.5") - 2*1440 margins = 9360 usable
                    # Center: 9360/2 = 4680; Right: 9360
                    tabs = pPr.find(qn('w:tabs'))
                    if tabs is None:
                        tabs = OxmlElement('w:tabs')
                        pPr.append(tabs)
                    else:
                        for old_tab in tabs.findall(qn('w:tab')):
                            tabs.remove(old_tab)
                    center_tab = OxmlElement('w:tab')
                    center_tab.set(qn('w:val'), 'center')
                    center_tab.set(qn('w:pos'), '4680')
                    center_tab.set(qn('w:leader'), 'none')
                    tabs.append(center_tab)
                    right_tab = OxmlElement('w:tab')
                    right_tab.set(qn('w:val'), 'right')
                    right_tab.set(qn('w:pos'), '9360')
                    right_tab.set(qn('w:leader'), 'none')
                    tabs.append(right_tab)

                    # Build leading tab run (moves equation to center)
                    lead_tab_r = OxmlElement('w:r')
                    lead_tab_el = OxmlElement('w:tab')
                    lead_tab_r.append(lead_tab_el)

                    # Build trailing tab run (moves to right for equation number)
                    trail_tab_r = OxmlElement('w:r')
                    trail_tab_el = OxmlElement('w:tab')
                    trail_tab_r.append(trail_tab_el)

                    # Build equation number run: "(N)"
                    eq_num_r = OxmlElement('w:r')
                    eq_rPr = OxmlElement('w:rPr')
                    eq_rFonts = OxmlElement('w:rFonts')
                    eq_rFonts.set(qn('w:ascii'), FONT_NAME)
                    eq_rFonts.set(qn('w:hAnsi'), FONT_NAME)
                    eq_rFonts.set(qn('w:eastAsia'), FONT_NAME)
                    eq_rFonts.set(qn('w:cs'), FONT_NAME)
                    eq_rPr.append(eq_rFonts)
                    eq_sz = OxmlElement('w:sz')
                    eq_sz.set(qn('w:val'), str(int(body_size * 2)))
                    eq_rPr.append(eq_sz)
                    eq_szCs = OxmlElement('w:szCs')
                    eq_szCs.set(qn('w:val'), str(int(body_size * 2)))
                    eq_rPr.append(eq_szCs)
                    eq_color = OxmlElement('w:color')
                    eq_color.set(qn('w:val'), '000000')
                    eq_rPr.append(eq_color)
                    eq_num_r.append(eq_rPr)
                    eq_t = OxmlElement('w:t')
                    eq_t.set(qn('xml:space'), 'preserve')
                    eq_t.text = f'({eq_counter})'
                    eq_num_r.append(eq_t)

                    # Unwrap oMath from oMathPara: extract oMath, insert tab-based layout
                    # parent is oMathPara; replace it with: lead_tab_r + oMath + trail_tab_r + eq_num_r
                    omath_para = parent  # m:oMathPara
                    omath_para.addprevious(lead_tab_r)
                    # Move oMath out of oMathPara to paragraph level
                    omath_para.addprevious(omath)
                    # Now omath is at paragraph level; add trailing elements after it
                    omath.addnext(eq_num_r)
                    omath.addnext(trail_tab_r)
                    # Remove the now-empty oMathPara
                    omath_para.getparent().remove(omath_para)

        # Clean w:instrText (field codes)
        for inst in p_elem.findall('.//' + qn('w:instrText')):
            if inst.text:
                if '~' in inst.text:
                    inst.text = replace_tildes_in_text(inst.text)
                if '\u2013' in inst.text or '\u2014' in inst.text:
                    inst.text = replace_dashes_in_text(inst.text)

    # --- Pass 2: Format all paragraphs ---
    before_abstract = False
    title_seen = False
    in_highlights = False  # Track highlight zone (between "Highlights" heading and Title)
    in_references = False  # Track references section (skip number-range → "to" conversion)

    for para in doc.paragraphs:
        style_name = (para.style.name if para.style else "").lower()
        text = para.text.strip()

        # Track highlight zone
        if text.lower() == "highlights" and is_heading(para):
            in_highlights = True
        elif is_title_paragraph(para):
            in_highlights = False

        # Track title/author zone
        if style_name == "title" or (not title_seen and is_heading(para) and text):
            title_seen = True
            before_abstract = True
        if style_name in ("abstract title", "abstract") or text.lower() == "abstract":
            before_abstract = False

        heading_level = is_heading(para)

        # Track references/acknowledgements section (once entered, stays true)
        # Skip number-range → "to" conversion in these sections
        if heading_level:
            section_text = text.lower().strip().lstrip('0123456789. ')
            if section_text in ('references', 'acknowledgements', 'acknowledgments',
                                'appendix', 'supplementary material'):
                in_references = True

        # Set line spacing and space after for all paragraphs
        para.paragraph_format.line_spacing = line_spacing
        space_after_pt = Pt(body_size * space_after) if space_after > 0 else Pt(0)
        para.paragraph_format.space_after = space_after_pt

        # --- Highlights items: add bullet prefix, clean text ---
        if in_highlights and not heading_level and text and text.lower() != "highlights":
            # Clean highlight text: remove relative improvements, reduce parens
            for run in para.runs:
                if run.text:
                    run.text = reduce_relative_improvements(run.text)
                    run.text = reduce_parentheses(run.text)
            format_highlight_bullet(para)
            # Auto-shorten highlights exceeding 85 chars
            if len(para.text.strip()) > 85:
                for run in para.runs:
                    if run.text:
                        # Common shortenings that preserve meaning
                        run.text = run.text.replace('district-level', 'district')
                        run.text = run.text.replace('run-to-run', 'run-to-run')  # keep
            set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.LEFT)
            para.paragraph_format.first_line_indent = Pt(0)
            for run in para.runs:
                set_run_font(run, size=body_size, bold=False, italic=False, color=FONT_COLOR)

        # --- Title: bold, centered, 14pt ---
        elif is_title_paragraph(para):
            set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.CENTER)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = space_after_pt
            for run in para.runs:
                set_run_font(run, size=14, bold=True, italic=False, color=FONT_COLOR)

        # --- Author/affiliation info: 11pt, centered ---
        elif is_author_info(para, before_abstract):
            set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.CENTER)
            para.paragraph_format.space_before = Pt(0)
            for run in para.runs:
                set_run_font(run, size=author_size, bold=False, italic=False, color=FONT_COLOR)

        # --- Abstract title: centered, bold ---
        elif style_name == "abstract title":
            set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.CENTER)
            para.paragraph_format.space_before = Pt(12)
            for run in para.runs:
                set_run_font(run, size=author_size, bold=True, italic=False, color=FONT_COLOR)

        # --- Abstract body: 11pt, justified, clean text ---
        elif style_name == "abstract":
            # Clean abstract text: remove relative improvements, reduce parens
            for run in para.runs:
                if run.text:
                    run.text = reduce_relative_improvements(run.text)
                    run.text = reduce_parentheses(run.text)
            set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.JUSTIFY)
            for run in para.runs:
                set_run_font(run, size=author_size, bold=False, italic=False, color=FONT_COLOR)

        # --- Keywords paragraph: bold label, fix spacing ---
        elif KEYWORDS_RE.match(text):
            set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.JUSTIFY)
            format_keywords_paragraph(para, body_size)

        # --- Equation paragraphs: preserve math, only set alignment & spacing ---
        # But if the paragraph also has Word numbering, treat as bullet (inline math in list item)
        elif is_equation_paragraph(para) and not has_word_numbering(para):
            for run in para.runs:
                set_run_font(run, size=body_size, bold=False, italic=False, color=FONT_COLOR)

        # --- Headings: bold, left-aligned ---
        elif heading_level:
            if needs_page_break_before(text):
                para.paragraph_format.page_break_before = True

            set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.LEFT)
            para.paragraph_format.space_before = Pt(12)
            for run in para.runs:
                set_run_font(run, size=body_size, bold=True, italic=False, color=FONT_COLOR)

        # --- Captions (by style or text): only label bold, no italic, centered ---
        elif is_caption_style(para) or is_caption(text):
            set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.CENTER)
            if is_caption(text):
                format_caption_runs(para, body_size)
            else:
                for run in para.runs:
                    set_run_font(run, size=body_size, bold=False, italic=False, color=FONT_COLOR)

        # --- Image paragraphs: centered ---
        elif is_image_paragraph(para):
            set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.CENTER)
            for run in para.runs:
                set_run_font(run, size=body_size, bold=False, italic=False, color=FONT_COLOR)

        # --- Body text: justified, no bold, no italic ---
        else:
            set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.JUSTIFY)
            # Bullet/list items: normalize to •, no first-line indent, justified
            if is_bullet_paragraph(para):
                normalize_bullet_symbol(para)
                para.paragraph_format.first_line_indent = Pt(0)
                set_paragraph_alignment(para, WD_ALIGN_PARAGRAPH.JUSTIFY)
            for run in para.runs:
                set_run_font(run, size=body_size, bold=False, italic=False, color=FONT_COLOR)

        # --- Replace dashes, tildes, math symbols, and orphan symbols in ALL paragraphs (except Title) ---
        if not is_title_paragraph(para):
            for run in para.runs:
                if run.text:
                    if '\u2013' in run.text or '\u2014' in run.text or re.search(r'\d\s*-\s*\d', run.text):
                        run.text = replace_dashes_in_text(run.text, skip_number_ranges=in_references)
                    if '~' in run.text or '∼' in run.text:
                        run.text = replace_tildes_in_text(run.text)
                    # Replace math symbols (≈, ≥, ≤, ±, <, >) with natural language
                    if any(c in run.text for c in '≈≥≤±') or re.search(r'[<>]\s*\d', run.text):
                        run.text = replace_math_symbols_in_text(run.text)
                    # Remove orphan math symbols (leftover from oMath flattening)
                    if run.text and run.text.strip() == '±':
                        run.text = ''

        # Remove trailing orphan math-symbol runs (oMath artifacts after main text)
        # Pattern: last substantial run ends with '.', ')', or '"'; subsequent short
        # runs contain only math symbols / Greek letters / operators — clear them.
        _MATH_ORPHAN_RE = re.compile(
            r'^[\s±≈∼≥≤≠∈μσεδαβγλ∑∏∫√∞NKDCM+\-*/=.,|(){}\[\]0-9]+$')
        runs = para.runs
        if runs:
            last_substantial = -1
            for ri in range(len(runs) - 1, -1, -1):
                if runs[ri].text and len(runs[ri].text.strip()) > 5:
                    last_substantial = ri
                    break
            if last_substantial >= 0:
                last_text = runs[last_substantial].text.rstrip()
                # Only clean if the main text ends with sentence-ending punctuation
                if last_text and last_text[-1] in '.;:)"\u201d':
                    for ri in range(last_substantial + 1, len(runs)):
                        rt = runs[ri].text
                        if rt and len(rt.strip()) <= 15 and _MATH_ORPHAN_RE.match(rt.strip()):
                            runs[ri].text = ''

    # Format tables (Times New Roman 11pt, three-line style)
    for table in doc.tables:
        format_table_fonts(table, size=table_size)

    # Update default style fonts
    for style in doc.styles:
        if hasattr(style, "font") and style.font is not None:
            style.font.name = FONT_NAME
            if hasattr(style, "element"):
                rPr = style.element.find(qn("w:rPr"))
                if rPr is not None:
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = rPr.makeelement(qn("w:rFonts"), {})
                        rPr.insert(0, rFonts)
                    rFonts.set(qn("w:ascii"), FONT_NAME)
                    rFonts.set(qn("w:hAnsi"), FONT_NAME)
                    rFonts.set(qn("w:eastAsia"), FONT_NAME)
                    rFonts.set(qn("w:cs"), FONT_NAME)

    save_path = output_path or docx_path
    doc.save(save_path)
    print(f"Formatted: {save_path}")
    print(f"  Style: Water Research / Elsevier")
    print(f"  Body font: {FONT_NAME}, {body_size}pt (black)")
    print(f"  Title: 14pt, bold")
    print(f"  Author/abstract: {author_size}pt, email → footnote")
    print(f"  Line spacing: {line_spacing}x")
    print(f"  Space after: {space_after} lines")
    print(f"  Margins: {margin} inch")
    print(f"  Headings: bold, left-aligned")
    print(f"  Captions: label bold only, no italic, centered")
    print(f"  Tables: three-line (三线表), full width, even columns, header bold")
    print(f"  Body: justified, no bold, dashes/tildes/math symbols replaced with natural language")
    print(f"  Equations: preserved")
    print(f"  References: page break before")


def main():
    parser = argparse.ArgumentParser(description="Format .docx for journal submission")
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("--output", help="Output path (default: overwrite input)")
    parser.add_argument("--body_size", type=float, default=11, help="Body font size in pt (default: 11)")
    parser.add_argument("--author_size", type=float, default=11, help="Author/abstract font size in pt (default: 11)")
    parser.add_argument("--table_size", type=float, default=11, help="Table font size in pt (default: 11)")
    parser.add_argument("--line_spacing", type=float, default=1.0, help="Line spacing (default: 1.0 single)")
    parser.add_argument("--space_after", type=float, default=0.5, help="Space after paragraphs in lines (default: 0.5)")
    parser.add_argument("--margin", type=float, default=1.0, help="Page margins in inches (default: 1.0)")
    args = parser.parse_args()

    format_document(
        args.input,
        output_path=args.output,
        body_size=args.body_size,
        author_size=args.author_size,
        table_size=args.table_size,
        line_spacing=args.line_spacing,
        space_after=args.space_after,
        margin=args.margin,
    )


if __name__ == "__main__":
    main()
